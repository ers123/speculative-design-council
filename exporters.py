"""Speculative Design Council — 산출물 내보내기 모듈.

원본: 연구 설계 도메인 exporters (지도교수/IRB/검색/BibTeX)
변환: SD 도메인 exporters
  - create_docx          : 종합 보고서 .docx (그대로 유지, 헤더만 SD 톤)
  - export_agents_md     : narratype 세계용 AGENTS.md (Stage 2 결과 기반)
  - export_constellation : STEEP 휠 좌표 + 4 archetype plot 데이터
  - export_diptych_seed  : Living Diptych 우측 패널 초기값
  - export_audience_report : 청중 반응 보고서 (Stage 3 종합)

설계 원칙 (그대로):
- 모든 변환기는 Ollama 로컬 호출 = $0 / 데이터 외부 안 나감
- 학생이 자기 최종 결과물에 즉시 활용 가능한 quality
"""

from __future__ import annotations

import io
import datetime
from typing import Callable

import ollama
from docx import Document


# ═════════════════════════════════════════════════════════════════════
# .docx (변경 없음 — 헤더는 templates.py의 DOCX_HEADER 사용)
# ═════════════════════════════════════════════════════════════════════

def create_docx(content: str, header: str, disclaimer: str) -> io.BytesIO:
    """Render Council 출력을 .docx로. 학생 최종 결과물 첨부용."""
    doc = Document()
    doc.add_heading(header, level=1)
    doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")

    p = doc.add_paragraph()
    run = p.add_run(disclaimer)
    run.italic = True
    doc.add_paragraph("")

    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("━"):
            doc.add_paragraph("─" * 40)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=2)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=3)
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═════════════════════════════════════════════════════════════════════
# Ollama-backed generators (그대로 유지)
# ═════════════════════════════════════════════════════════════════════

_EMPTY_CONTENT_HINT = (
    "[빈 응답 — 사고 모드가 켜진 상태에서 예산 내에 최종 응답을 내지 못했을 수 있습니다. "
    "사고 모드를 끄거나 더 작은 산출물부터 시도하세요.]"
)


def _ollama_once(model: str, system_prompt: str, user_payload: str, max_tokens: int = 2048,
                 think: bool = False, keep_alive: int = 0) -> str:
    """Single non-streaming ollama.chat call. think=False / keep_alive=0 (VRAM hygiene)."""
    try:
        response = ollama.chat(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_payload},
            ],
            think=think,
            options={"temperature": 0.2, "top_p": 0.9, "num_predict": max_tokens},
            keep_alive=keep_alive,
            stream=False,
        )
        msg = response.get("message") if isinstance(response, dict) else getattr(response, "message", None)
        if msg is None:
            return _EMPTY_CONTENT_HINT
        content = msg.get("content", "") if isinstance(msg, dict) else (getattr(msg, "content", "") or "")
        if not content.strip():
            return _EMPTY_CONTENT_HINT
        return content
    except Exception as e:
        return f"[변환 실패: {e}]"


# ═════════════════════════════════════════════════════════════════════
# 1. AGENTS.md — narratype 세계용 protocol 문서
# ═════════════════════════════════════════════════════════════════════

_AGENTS_MD_PROMPT = """당신은 학생의 narratype 세계용 AGENTS.md를 최종 형식으로 정리합니다.

agents.md (Antigravity·Claude Code·Codex 협업 표준) 의 markdown 문법을 정확히 따르세요.
출력은 학생이 자기 spec design artifact(웹사이트·인쇄물·전시 자료)에 그대로 첨부 가능한 quality여야 합니다.

출력 구조:

# AGENTS.md — [narratype 세계명]
> 이 문서는 [세계명] 안에서 활동하는 AI 에이전트들의 행동 강령이다.
> 작성: [학생명] / 시기: 2045 / 버전: 1.0

## 1. 세계 환경
[2-3 문장 세계 묘사]

## 2. 에이전트 명단
### 2.1 [이름] (분류)
- 역할:
- 권한 수준: L1/L2/L3
- 통신 언어:
- 신뢰 등급:
- 시민 거부권:

(2.2 ~ 2.5 동일 형식 — 최소 3명, 권장 5명)

## 3. 협업 프로토콜
### 3.1 우선순위 매트릭스
[Markdown 표]

### 3.2 충돌 해결
[3-5단계 절차]

### 3.3 검열·차단 정책
[명시적 규칙]

### 3.4 기록·감사
[로그 형식, 보관 기간, 공개 범위]

## 4. 시민 인터페이스
[시민이 에이전트와 만나는 방식, 거부권, 항소 절차]

## 5. 책임 소재
[에이전트 결정의 책임은 누구에게]

## 6. 변경 이력
- 2045-XX-XX: 초기 작성

---
이 AGENTS.md는 [학생 narratype 작품]의 일부이다.

출력은 ```markdown 펜스 없이 순수 markdown으로 시작하세요."""


def export_agents_md(model: str, director_output: str, student_name: str = "학생") -> str:
    payload = (
        f"=== Council 종합 결과 (Stage 2: AGENTS.md 작성) ===\n{director_output}\n=== 끝 ===\n\n"
        f"학생명: {student_name}\n"
        f"이것을 위 형식의 최종 AGENTS.md로 정리하세요. 작가의 결정을 존중하되 markdown 문법은 정확히."
    )
    return _ollama_once(model, _AGENTS_MD_PROMPT, payload, max_tokens=2500)


# ═════════════════════════════════════════════════════════════════════
# 2. Constellation 좌표 — STEEP 휠 + 4 archetype
# ═════════════════════════════════════════════════════════════════════

_CONSTELLATION_PROMPT = """당신은 학생의 narratype을 STEEP 휠 위 좌표로 변환합니다.

STEEP 휠 4사분면:
- Social (사회)
- Cultural (문화)
- Environmental (환경)
- Technological (기술)

축:
- 안쪽(중심) → 바깥(주변): Now → Future (1=지금, 5=먼 미래)
- 4 archetype별로 별 1개씩 plot

출력 형식 (.txt, 학생이 Miro 또는 인쇄물에 그대로 옮길 수 있는 좌표 데이터):

# Constellation 좌표 — [narratype 세계명]
# Generated: YYYY-MM-DD HH:MM

## STEEP 4사분면 좌표 (각 archetype별)
# format: archetype | quadrant | distance_from_center (1-5) | label_short | label_long

성장   | technological | 4 | "기술 성장" | "narratype의 기술-성장 시나리오 한 줄"
붕괴   | social        | 5 | "사회 붕괴" | ...
지속   | cultural      | 3 | "문화 지속" | ...
변혁   | environmental | 5 | "환경 변혁" | ...

## 학생 본인의 두 별 (워크숍 50-60분 단계)
거울_AI | technological | 2 | "거울 AI" | "폰의 Gemma 4 — 학생 narratype의 어떤 측면을 비춤"
지도_AI | social        | 3 | "지도 AI" | "노트북 cloud — 학생 narratype의 어떤 측면을 비춤"

## 잇는 선 (= AGENTS.md protocol)
거울_AI <-> 지도_AI : "AGENTS.md 핵심 규칙 1줄"

## Constellation의 이름 (이 별자리의 mythology)
[3-5단어 별자리 이름 + 1줄 설명]

---
원칙:
- 각 archetype 좌표는 narratype의 강점에 따라 거리·사분면 다르게
- 4 archetype이 서로 다른 사분면에 위치할 수 있음 (다양성)
- mythology 이름은 한국어, 시적으로"""


def export_constellation(model: str, all_directors_output: str) -> str:
    payload = (
        f"=== 모든 Stage Director 출력 ===\n{all_directors_output}\n=== 끝 ===\n\n"
        f"위 분석을 STEEP 휠 좌표 데이터로 변환하세요."
    )
    return _ollama_once(model, _CONSTELLATION_PROMPT, payload, max_tokens=1500)


# ═════════════════════════════════════════════════════════════════════
# 3. Living Diptych 시드 — Alt 1 결과물 우측 패널 초기값
# ═════════════════════════════════════════════════════════════════════

_DIPTYCH_SEED_PROMPT = """당신은 학생이 최종 결과물 옵션 1 (Living Diptych)을 만들 때 우측 패널의 초기 대화 시드를 작성합니다.

Living Diptych 형식:
- 좌측: 학생의 narratype artifact (정적)
- 우측: 두 AI(거울·지도)의 끝없는 대화 (자라남)

이 출력은 우측 패널의 첫 5-10개 turn을 미리 작성한 것입니다.
학생의 사이트가 처음 열렸을 때 보일 대화입니다. 이후 turn은 학생의 AGENTS.md protocol에 따라 자동 생성됨.

출력 형식 (.md, 학생이 자기 사이트의 첫 콘텐츠로 사용):

# Living Diptych — 우측 패널 시드
# narratype: [세계명]

---

### Turn 1 — 거울 AI (YYYY-MM-DD)
[이 작품을 처음 본 거울 AI의 응답, 4-6문장, 1인칭]

### Turn 2 — 지도 AI (YYYY-MM-DD)
[거울 AI에 대한 반응, 4-6문장, 비판적 톤]

### Turn 3 — 거울 AI
[반박 또는 보완, 3-4문장]

### Turn 4 — 지도 AI
[심화 질문, 3-4문장]

### Turn 5 — 거울 AI
[학생 작품의 어떤 element를 다시 봐야 하는지 제안]

(Turn 6-10도 같은 형식, 시간 흐름 표현)

---
이후 자동 생성: AGENTS.md protocol 적용.

원칙:
- 두 AI의 인격이 명확히 구분되어야 함 (거울=확장적, 지도=비판적)
- 학생의 narratype을 배경으로 하되 추상에 머물지 않음
- 시간 표현 (날짜·시간) 포함
- 한국어"""


def export_diptych_seed(model: str, all_directors_output: str) -> str:
    payload = (
        f"=== 모든 Stage Director 출력 ===\n{all_directors_output}\n=== 끝 ===\n\n"
        f"위 분석을 바탕으로 Living Diptych 우측 패널의 첫 5-10 turn 작성."
    )
    return _ollama_once(model, _DIPTYCH_SEED_PROMPT, payload, max_tokens=2500)


# ═════════════════════════════════════════════════════════════════════
# 4. 청중 반응 보고서 — Stage 3 종합
# ═════════════════════════════════════════════════════════════════════

_AUDIENCE_REPORT_PROMPT = """당신은 학생이 최종 결과물에 첨부할 "청중 반응 분석 보고서"를 정리합니다.
2025년의 한계 — "친구·가족 중심 동질적 반응" — 를 명시적으로 넘어선 evidence-based 분석.

출력 형식 (.md, 학생 작품 부록 또는 별첨 자료로 활용):

# 청중 반응 분석 보고서 — [narratype 세계명]
# 페르소나: Nemotron-Personas-Korea (NVIDIA, CC-BY-4.0)
# Generated: YYYY-MM-DD HH:MM

## 1. Executive Summary
[3-4 문장 — 청중이 작품을 어떻게 받았는지, 가장 두드러진 패턴, 학생 주목 포인트]

## 2. 4 Archetype 분포

```
성장 ████████░░ 40%
붕괴 ██████░░░░ 30%
지속 ████░░░░░░ 20%
변혁 ██░░░░░░░░ 10%
```

(각 분포의 demographic 패턴 1-2문장)

## 3. 인용할 페르소나 반응 5선

### 3.1 가장 강력한 찬성
> [페르소나 인용 1-2문장]
출처: [페르소나 이름·나이·지역·직업]

### 3.2 가장 강력한 반대
> [인용]
출처: [...]

### 3.3 가장 의외였던 반응
> [인용]
출처: [...]

### 3.4 가장 정확한 비평
> [인용]
출처: [...]

### 3.5 가장 무서운 한 마디
> [인용]
출처: [...]

## 4. 학생 작품의 demographic 사각지대

| demographic | 작품 안에서의 위치 | 사라진 이유 (추정) |
|-------------|-------------------|-------------------|
| ... | ... | ... |

## 5. 다음에 답해야 할 질문 3개

1. ...
2. ...
3. ...

---
주의: 이 페르소나 반응은 합성 시뮬레이션이며, 실제 fieldwork(대중 반응 분석)의 보완재이지 대체재가 아닙니다. 응답 패턴 자체를 비판적으로 분석하는 것이 곧 SD 작업입니다."""


def export_audience_report(model: str, director_output: str) -> str:
    payload = f"=== Stage 3 (청중 반응) Director 출력 ===\n{director_output}\n=== 끝 ==="
    return _ollama_once(model, _AUDIENCE_REPORT_PROMPT, payload, max_tokens=2500)


# ═════════════════════════════════════════════════════════════════════
# Export 디스패치 (UI에서 참조)
# ═════════════════════════════════════════════════════════════════════

EXPORT_REGISTRY = {
    "docx": {
        "label": "Word 보고서 (.docx) — 최종 결과물 첨부용",
        "filename": "SD_Council_보고서",
        "ext": "docx",
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "kind": "docx",
    },
    "agents_md": {
        "label": "AGENTS.md (.md) — narratype 세계 protocol 문서",
        "filename": "AGENTS",
        "ext": "md",
        "mime": "text/markdown",
        "kind": "llm",
        "fn": export_agents_md,
    },
    "constellation": {
        "label": "Constellation 좌표 (.txt) — STEEP 휠 + 4 archetype plot",
        "filename": "constellation_coords",
        "ext": "txt",
        "mime": "text/plain",
        "kind": "llm",
        "fn": export_constellation,
    },
    "diptych_seed": {
        "label": "Living Diptych 시드 (.md) — Alt 1 결과물 우측 패널",
        "filename": "diptych_seed",
        "ext": "md",
        "mime": "text/markdown",
        "kind": "llm",
        "fn": export_diptych_seed,
    },
    "audience_report": {
        "label": "청중 반응 보고서 (.md) — 페르소나 분석",
        "filename": "audience_report",
        "ext": "md",
        "mime": "text/markdown",
        "kind": "llm",
        "fn": export_audience_report,
    },
}
